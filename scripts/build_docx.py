#!/usr/bin/env python3
"""Build forprosjekt.docx for the HMS school assignment.

Run from project root:
    python3 scripts/build_docx.py

Produces: forprosjekt.docx in the project root.
"""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

ACCENT = RGBColor(0xC2, 0x41, 0x0C)
INK = RGBColor(0x1C, 0x1A, 0x17)
MUTED = RGBColor(0x5A, 0x54, 0x4C)


def set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = INK if level > 1 else ACCENT


def add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------


def build_forside(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nForprosjekt")
    run.font.size = Pt(28)
    run.font.color.rgb = MUTED
    run.bold = False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HMS — Arbeid ved dataskjerm")
    run.font.size = Pt(40)
    run.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Et nettsted laget for jevnaldrende elever.")
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        ("Navn:", "[NAVN]"),
        ("Klasse:", "Vg2 Informasjonsteknologi og medieproduksjon"),
        ("Fag:", "IT-utvikling 1 / Mediedesign og medieuttrykk"),
        ("Levert:", "[DATO]"),
    ]
    table = doc.add_table(rows=len(info_lines), cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (k, v) in enumerate(info_lines):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_page_break()


def build_innledning(doc: Document) -> None:
    add_heading(doc, "1. Innledning", level=1)
    add_paragraph(
        doc,
        "Oppgaven gikk ut på å lage et nettsted som forklarer HMS ved arbeid ved "
        "dataskjerm — særlig myntet på elever og unge voksne som akkurat har begynt "
        "å sitte mye foran skjerm. Nettstedet skal dekke risiko for helseplager, "
        "krav til fysisk utforming av arbeidsplassen, og enkle øvelser man kan ta "
        "i hverdagen.",
    )
    add_paragraph(
        doc,
        "Forprosjektet beskriver hvordan jeg har planlagt arbeidet, hvilke verktøy "
        "jeg har vurdert å bruke, og hvorfor jeg landet på å bygge nettstedet for "
        "hånd med ren HTML, CSS og litt JavaScript.",
    )


def build_skisse(doc: Document) -> None:
    add_heading(doc, "2. Skisse av nettstedet", level=1)
    add_paragraph(
        doc,
        "Nettstedet er bygget rundt et editorial / magasin-aktig oppsett: ett "
        "dominerende bilde-/illustrasjonsfelt på forsiden, deretter klare seksjoner "
        "med korte tekstavsnitt og handlingsorienterte kort som leder videre. "
        "Designet er bevisst varmere enn typiske myndighetssider — terrakottafarge, "
        "Fraunces som overskriftsfont og Inter som brødtekst.",
    )
    add_heading(doc, "ASCII-wireframe — forside", level=2)
    wire = (
        "+----------------------------------------------------------+\n"
        "|  [LOGO]  HMS ved skjerm        Forside Risiko Ergonomi…  |\n"
        "+----------------------------------------------------------+\n"
        "|                                                          |\n"
        "|  Skjermarbeid UTEN verkende skuldre.                     |\n"
        "|  Kort lead-tekst om hvem siden er for.                   |\n"
        "|  [Start med risiko →]   [Hopp til ergonomi]              |\n"
        "|                                                          |\n"
        "|  +--- Illustrasjon (skrivebord, lampe, kaffe) ---+       |\n"
        "|                                                          |\n"
        "+----------------------------------------------------------+\n"
        "|  Hva er HMS ved dataskjerm? (kort prosa)                 |\n"
        "+----------------------------------------------------------+\n"
        "|  Risiko: 3 kort   [Muskel] [Syn] [Psyk]                  |\n"
        "+----------------------------------------------------------+\n"
        "|  Tall som gjør deg våken: 46% / 20% / 15%                |\n"
        "+----------------------------------------------------------+\n"
        "|  Ergonomi: 4 kort   [Bord] [Skjerm] [Tast] [Mus]         |\n"
        "+----------------------------------------------------------+\n"
        "|  CTA: «Tre minutter, og du kan resten av livet.»         |\n"
        "+----------------------------------------------------------+\n"
        "|  Footer: om · innhold · hovedkilder                      |\n"
        "+----------------------------------------------------------+\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(wire)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def build_flytskjema(doc: Document) -> None:
    add_heading(doc, "3. Flytskjema — menystruktur", level=1)
    add_paragraph(
        doc,
        "Nettstedet har to nivåer: ti-elements toppmeny med to nedtrekksmenyer, "
        "og 11 sider totalt. Tabellen viser hierarkiet og hvilke sider som lenker "
        "til hvilke.",
    )

    pages = [
        ("Forside", "index.html", "Hero + intro + lenker til Risiko og Ergonomi"),
        ("Risiko (oversikt)", "risiko/index.html", "3 kort til underseksjonene"),
        ("  Muskelplager", "risiko/muskelplager.html", "Statistikk, tiltak, 3 øvelser"),
        ("  Syn", "risiko/syn.html", "20-20-20-regelen, databriller, lys"),
        ("  Psykososialt", "risiko/psykososialt.html", "STAMI 15 %, krav fra 2026"),
        ("Ergonomi (oversikt)", "ergonomi/index.html", "4 kort til underseksjonene"),
        ("  Arbeidsbord", "ergonomi/arbeidsbord.html", "Matt overflate, hev/senk, mål"),
        ("  Skjerm", "ergonomi/skjerm.html", "Høyde, avstand, lys, anti-refleks"),
        ("  Tastatur", "ergonomi/tastatur.html", "Separat, lav profil, ergonomisk"),
        ("  Datamus", "ergonomi/datamus.html", "Plassering, vertikal, trackball"),
        ("Kilder", "kilder.html", "Annoterte og grupperte kilder"),
        ("Om", "om.html", "Refleksjon CMS/KI vs hardkoding"),
    ]
    table = doc.add_table(rows=1 + len(pages), cols=3)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Side"
    hdr[1].text = "Filsti"
    hdr[2].text = "Innhold"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
        set_cell_shading(cell, "FCEAD8")
    for i, (name, path, desc) in enumerate(pages, start=1):
        row = table.rows[i].cells
        row[0].text = name
        row[1].text = path
        row[2].text = desc

    add_paragraph(
        doc,
        "I tillegg: hver side har breadcrumbs øverst og en \"På denne siden\"-meny "
        "(TOC) i sidefeltet på lange artikler.",
        italic=True,
    )


def build_cms_ki(doc: Document) -> None:
    add_heading(doc, "4. Valg av verktøy: hardkoding, CMS eller KI?", level=1)
    add_paragraph(
        doc,
        "Jeg vurderte tre tilnærminger før jeg startet: hardkoding (ren HTML/CSS/JS), "
        "et CMS som WordPress eller Wix, og et KI-verktøy som v0, Lovable eller Framer. "
        "Hver av dem løser oppgaven, men på veldig forskjellig måte.",
    )
    add_heading(doc, "Hva er hva?", level=2)
    add_bullets(
        doc,
        [
            "Hardkoding: man skriver alt selv. Tar lengst tid, gir mest kontroll, "
            "lærer mest.",
            "CMS (Content Management System): WordPress, Wix, Squarespace, Webflow. "
            "Du velger mal, drar inn blokker, redigerer tekst.",
            "KI-verktøy: v0 (av Vercel), Lovable, Framer AI, Bolt. Du beskriver hva "
            "du vil ha, KI-en genererer koden.",
        ],
    )
    add_heading(doc, "Hvorfor jeg valgte hardkoding", level=2)
    add_numbered(
        doc,
        [
            "Læringsutbytte. Linjen heter Informasjonsteknologi og medieproduksjon "
            "— det er litt poenget å forstå hvordan en nettside er bygd.",
            "Full kontroll over design og kode. Et CMS-tema fører deg ofte i "
            "bestemte retninger jeg ikke vil.",
            "Ingen lock-in. Filene fungerer offline og kan flyttes hvor som helst.",
            "Gratis. Ingen abonnement, ingen domeneavhengighet for selve koden.",
        ],
    )


def build_sammenligning(doc: Document) -> None:
    add_heading(doc, "5. Sammenligning: hardkoding vs. CMS/KI-verktøy", level=1)

    rows = [
        (
            "Hardkoding (HTML/CSS/JS)",
            [
                "Full kontroll over kode, design og ytelse",
                "Gratis — ingen abonnement",
                "Ingen lock-in",
                "Best for læring av webteknologier",
            ],
            [
                "Tar lengre tid å sette opp",
                "Vanskelig å oppdatere innhold for ikke-utviklere",
                "Du må selv håndtere tilgjengelighet og responsivt design",
            ],
        ),
        (
            "CMS (WordPress, Wix, Squarespace)",
            [
                "Rask oppstart — fungerende side på timer",
                "Innholdsforvaltning for ikke-utviklere",
                "Ferdige temaer dekker tilgjengelighet og responsivt design",
            ],
            [
                "Lock-in — vanskelig å bytte plattform",
                "Abonnementskostnader",
                "Mindre kontroll over kode og SEO",
                "Tendens til at \"alle sider ser like ut\"",
            ],
        ),
        (
            "KI-verktøy (v0, Lovable, Framer AI)",
            [
                "Ekstremt rask — utkast på minutter",
                "Bra som idégenerator selv om du hardkoder",
                "Sparker i gang \"hvordan skal det se ut\"-fasen",
            ],
            [
                "Resultatet ofte generisk \"AI-glass-og-gradient\"-stil",
                "Mister forståelsen av hvorfor koden er som den er",
                "Mange feil i tilgjengelighet og semantisk markup",
                "Krever at du forstår koden — ikke trygt å bruke blindt",
            ],
        ),
    ]

    for tilnærming, fordeler, ulemper in rows:
        add_heading(doc, tilnærming, level=2)

        p = doc.add_paragraph()
        r = p.add_run("Fordeler")
        r.bold = True
        r.font.color.rgb = ACCENT
        add_bullets(doc, fordeler)

        p = doc.add_paragraph()
        r = p.add_run("Ulemper")
        r.bold = True
        r.font.color.rgb = ACCENT
        add_bullets(doc, ulemper)

    add_heading(doc, "Konklusjonen for dette prosjektet", level=2)
    add_paragraph(
        doc,
        "Jeg landet på hardkoding fordi læringsutbytte var det viktigste målet. "
        "KI ble brukt som sparringspartner for design og formuleringer — ikke som "
        "kodegenerator. All HTML, CSS og JavaScript er skrevet, lest og forstått "
        "av meg.",
    )


def build_tidsplan(doc: Document) -> None:
    add_heading(doc, "6. Tidsplan", level=1)
    add_paragraph(
        doc,
        "Plan for periodene fra uke 14 til 17. Tabellen viser hva jeg har planlagt "
        "å gjøre per uke og estimert tid.",
    )
    rows = [
        ("Uke", "Hva", "Mål", "Estimat"),
        ("14", "Forprosjekt + research", "Levere forprosjekt-dokumentet, kildeliste klar", "8 t"),
        ("15", "Designsystem + forside + risiko-sider", "Forside og 4 risiko-sider ferdig", "12 t"),
        ("16", "Ergonomi-sider + kilder/om + JavaScript", "Alle 11 sider live, app.js + tema-toggle", "12 t"),
        ("17", "Egne bilder/øvelser, polish, innlevering", "Bilder lagt inn, alle TODO-er løst, levert", "8 t"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Light Grid"
    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = val
            if i == 0:
                for run in cells[j].paragraphs[0].runs:
                    run.bold = True
                set_cell_shading(cells[j], "FCEAD8")


def build_logg_mal(doc: Document) -> None:
    add_heading(doc, "7. Logg-mal", level=1)
    add_paragraph(
        doc,
        "Tom mal for arbeidslogg — fyll inn etter hver økt. Jeg fyller inn det "
        "som faktisk er gjort, ikke det jeg hadde tenkt å gjøre.",
    )
    rows = [["Dato", "Hva ble gjort", "Tid (t)", "Kommentar / problemer"]]
    for _ in range(10):
        rows.append(["", "", "", ""])
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Light Grid"
    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = val
            if i == 0:
                for run in cells[j].paragraphs[0].runs:
                    run.bold = True
                set_cell_shading(cells[j], "FCEAD8")


def build_sjekkliste(doc: Document) -> None:
    add_heading(doc, "8. Sjekkliste — innleveringskrav", level=1)
    items = [
        "Forside med tydelig verdiforslag og kort introduksjon til HMS",
        "Egen seksjon for risiko, med tre underseksjoner (muskel, syn, psyk)",
        "Egen seksjon for ergonomi, med fire underseksjoner (bord, skjerm, tast, mus)",
        "Tre øvelser med eget bilde/video (lagt inn av meg, ikke plassholder)",
        "Egen kilder-side med alle kilder lenket og kort forklart",
        "Egen om-side med refleksjon rundt verktøyvalg",
        "Sticky topp-meny med dropdowns på desktop og hamburger på mobil",
        "Brødsmuler på alle underseksjoner",
        "På denne siden-meny på lange artikler",
        "Lyst og mørkt tema med manuelt valg som lagres",
        "Tilgjengelighet: tab-rekkefølge OK, fokusringer synlige, alt-tekst på bilder",
        "Responsivt: testet på 375 / 768 / 1024 / 1280 / 1440 piksler",
        "Alle faktautsagn er paraframert og lenker til kilde",
        "Forprosjekt-dokumentet (denne filen) er fylt ut og levert",
    ]
    for item in items:
        p = doc.add_paragraph()
        r = p.add_run("☐  ")
        r.font.size = Pt(13)
        r2 = p.add_run(item)
        r2.font.size = Pt(11)


def build_kildeliste(doc: Document) -> None:
    add_heading(doc, "9. Kildeliste", level=1)
    add_paragraph(
        doc,
        "Alle kilder brukt på nettstedet, gruppert etter type. Sidetallene viser "
        "til hvor de oppstår på nettstedet.",
    )

    groups = [
        (
            "Lover og forskrifter",
            [
                ("Forskrift om arbeid ved dataskjerm (best.nr. 528)",
                 "https://lovdata.no/dokument/SF/forskrift/2011-12-06-1356/"),
                ("Arbeidsmiljøloven",
                 "https://lovdata.no/dokument/NL/lov/2005-06-17-62"),
            ],
        ),
        (
            "Arbeidstilsynet",
            [
                ("Arbeid ved dataskjerm",
                 "https://www.arbeidstilsynet.no/tema/ergonomi/arbeid-ved-dataskjerm/"),
                ("Risiko for helseplager ved arbeid ved dataskjerm",
                 "https://www.arbeidstilsynet.no/tema/ergonomi/arbeid-ved-dataskjerm/risiko-for-helseplager-ved-arbeid-ved-dataskjerm/"),
                ("Synsundersøking og databriller",
                 "https://www.arbeidstilsynet.no/tema/ergonomi/arbeid-ved-dataskjerm/synsundersokelse-og-databriller/"),
                ("Psykososialt arbeidsmiljø",
                 "https://www.arbeidstilsynet.no/tema/psykososialt-arbeidsmiljo/"),
            ],
        ),
        (
            "Forskning og statistikk",
            [
                ("STAMI — Statens arbeidsmiljøinstitutt", "https://www.stami.no/"),
                ("Faktabok om arbeidsmiljø og helse", "https://www.stami.no/faktabok-om-arbeidsmiljo-og-helse/"),
                ("Helsenorge — Muskel og skjelett", "https://www.helsenorge.no/sykdom/muskel-og-skjelett/"),
            ],
        ),
        (
            "Bedriftshelsetjeneste / praktiske veiledere",
            [
                ("Avonova", "https://www.avonova.no/"),
                ("Norsk Fysioterapeutforbund", "https://www.fysio.no/"),
                ("AJ Produkter — Ergonomi", "https://www.ajprodukter.no/inspirasjon/ergonomi/"),
            ],
        ),
        (
            "Produktreferanser",
            [
                ("AJ Produkter Flexus (hev/senk-bord)", "https://www.ajprodukter.no/"),
                ("IKEA Bekant (hev/senk-bord)", "https://www.ikea.com/no/no/p/bekant-skrivebord-sitt-staa-hvit-s89022527/"),
                ("Logitech Ergo K860 (delt tastatur)", "https://www.logitech.com/no-no/products/keyboards/ergo-k860-split-keyboard.html"),
                ("Microsoft Sculpt Ergonomic", "https://www.microsoft.com/en-us/d/microsoft-sculpt-ergonomic-keyboard/8r2hw03gpj4j"),
                ("Logitech MX Vertical (vertikal mus)", "https://www.logitech.com/no-no/products/mice/mx-vertical-ergonomic-mouse.html"),
            ],
        ),
    ]

    for title, items in groups:
        add_heading(doc, title, level=2)
        for label, url in items:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(label + " — ")
            r.font.size = Pt(11)
            r2 = p.add_run(url)
            r2.font.size = Pt(10)
            r2.font.color.rgb = ACCENT


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    build_forside(doc)
    build_innledning(doc)
    build_skisse(doc)
    build_flytskjema(doc)
    build_cms_ki(doc)
    build_sammenligning(doc)
    build_tidsplan(doc)
    build_logg_mal(doc)
    build_sjekkliste(doc)
    build_kildeliste(doc)

    output = os.path.join(os.path.dirname(__file__), "..", "forprosjekt.docx")
    output = os.path.abspath(output)
    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
